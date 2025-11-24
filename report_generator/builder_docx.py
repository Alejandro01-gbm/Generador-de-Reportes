from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FIELD_ORDER = [
    "No de alerta", "Criticidad", "Reportado por", "Descripción de la alerta",
    "Fecha y hora de Inicio de la alerta", "Total de Eventos", "Fuentes de Logs",
    "IP Origen", "IP Destino", "Evento contenido",
    "Indicadores de Compromiso (IoCs)", "Cuenta/s", "Análisis", "Recomendaciones",
]

def _style_doc(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(12)

def _available_width_inches(doc: Document) -> float:
    section = doc.sections[0]
    emu_to_inch = 1 / 914400
    return (section.page_width - section.left_margin - section.right_margin) * emu_to_inch

def build_docx(data: dict, outfile: str):
    doc = Document()
    _style_doc(doc)

    # Márgenes equilibrados
    section = doc.sections[0]
    section.left_margin   = Inches(0.7)
    section.right_margin  = Inches(0.7)
    section.top_margin    = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    # === Tabla principal (2 columnas) ===
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Fila 0: encabezado con imagen (merge)
    hdr_cell = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    hdr_para = hdr_cell.paragraphs[0]
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    img_path = os.path.join(os.getcwd(), "img", "Notificacion_de_seguridad.jpg")
    if os.path.exists(img_path):
        run = hdr_para.add_run()
        run.add_picture(img_path, width=Inches(_available_width_inches(doc)))
    else:
        r = hdr_para.add_run("NOTIFICACIÓN DE SEGURIDAD")
        r.bold = True
        r.font.size = Pt(16)

    # === Celdas Campo | Valor ===
    for field in FIELD_ORDER:
        row = table.add_row().cells

        # Columna izquierda:
        left = row[0]
        left.text = "" 
        left_run = left.paragraphs[0].add_run(field + ":")
        left_run.bold = True

        # Columna derecha: 
        if field == "Indicadores de Compromiso (IoCs)":
            raw_text = str(data.get(field, "") or "")
            paragraph = row[1].paragraphs[0]
            for line in raw_text.split("\n"):
                txt = (line or "").strip()
                if not txt:
                    paragraph.add_run("\n")
                    continue
                if txt.lower().startswith(("malware name", "hash", "ip maliciosa", "puertos")):
                    r = paragraph.add_run(txt + "\n")
                    r.bold = True
                else:
                    paragraph.add_run(txt + "\n")
        else:
            row[1].text = str(data.get(field, ""))

    doc.add_paragraph("")
    doc.save(outfile)
    print(f"[OK] Reporte generado en: {outfile}")
